import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
from docx.oxml.ns import qn

def dump(fn):
    d = docx.Document(fn)
    print('\n===== FILE:', fn, '=====')
    body = d.element.body
    tbl_idx = 0
    para_map = {p._p: p for p in d.paragraphs}
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = para_map.get(child)
            if p is not None and p.text.strip():
                print(p.text)
        elif child.tag == qn('w:tbl'):
            t = d.tables[tbl_idx]; tbl_idx += 1
            print('--- TABLE ---')
            for row in t.rows:
                print(' | '.join(c.text.strip() for c in row.cells))
            print('--- END TABLE ---')

for fn in sys.argv[1:]:
    dump(fn)
